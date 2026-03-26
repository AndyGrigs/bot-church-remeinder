import calendar
from telegram import Update, ReplyKeyboardMarkup, KeyboardButton
from telegram.ext import Application, CommandHandler, ContextTypes, MessageHandler, filters
from dotenv import load_dotenv
import os
from datetime import datetime, time
import docx  # python-docx
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches


# Завантаження змінних із .env файлу
load_dotenv()
BOT_TOKEN = os.getenv('BOT_TOKEN')
ADMIN_CHAT_ID = os.getenv('ADMIN_CHAT_ID')   # група де вводять дані
GROUP_CHAT_ID = os.getenv('GROUP_CHAT_ID')   # група куди йдуть нагадування

if not BOT_TOKEN or not ADMIN_CHAT_ID or not GROUP_CHAT_ID:
    print("Помилка: BOT_TOKEN, ADMIN_CHAT_ID або GROUP_CHAT_ID не встановлено. Перевірте файл .env.")
    exit(1)

try:
    ADMIN_CHAT_ID = int(ADMIN_CHAT_ID)
    GROUP_CHAT_ID = int(GROUP_CHAT_ID)
except ValueError:
    print("Помилка: ADMIN_CHAT_ID і GROUP_CHAT_ID повинні бути цілими числами.")
    exit(1)


def is_admin_chat(update: Update) -> bool:
    """Перевіряє, що команда надійшла з адмін-групи."""
    return update.effective_chat.id == ADMIN_CHAT_ID

user_states = {}

SCHEDULE_FILE = "schedule.txt"
EVENTS_FILE = "events.txt"

# Функція для встановлення кольору фону (заливки) клітинки:
def set_cell_bg_color(cell, color_hex: str):
    """
    Заливка клітинки таблиці коліром у форматі HEX (без #), наприклад "FF0000" (червоний).
    """
    # Отримуємо безпосередньо <w:tc> XML-елемент
    tc_pr = cell._tc.get_or_add_tcPr()
    # Створюємо <w:shd w:fill="COLOR_HEX"/>
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)

def load_schedule():
    """Завантаження розкладу з txt файлу."""
    schedule = {}
    if not os.path.exists(SCHEDULE_FILE):
        return schedule
    with open(SCHEDULE_FILE, "r", encoding="utf-8") as f:
        for line in f:
            line = line.strip()
            if not line:
                continue
            date, preachers_str = line.split("|", 1)
            schedule[date] = preachers_str.split(",")
    return schedule

def save_schedule_to_file(schedule: dict):
    """Записує весь розклад у txt файл."""
    with open(SCHEDULE_FILE, "w", encoding="utf-8") as f:
        for date, preachers in schedule.items():
            f.write(f"{date}|{','.join(preachers)}\n")

def load_events():
    """Завантаження подій з txt файлу."""
    events = []
    if not os.path.exists(EVENTS_FILE):
        return events
    with open(EVENTS_FILE, "r", encoding="utf-8") as f:
        for line in f:
            line = line.strip()
            if not line:
                continue
            date, title = line.split("|", 1)
            events.append({"date": date, "title": title})
    return events

def save_event(date: str, title: str):
    """Збереження події в txt файл."""
    with open(EVENTS_FILE, "a", encoding="utf-8") as f:
        f.write(f"{date}|{title}\n")

def delete_event(date: str, title: str) -> bool:
    """Видалення події з txt файлу."""
    events = load_events()
    new_events = [e for e in events if not (e["date"] == date and e["title"] == title)]
    if len(new_events) == len(events):
        return False
    with open(EVENTS_FILE, "w", encoding="utf-8") as f:
        for e in new_events:
            f.write(f"{e['date']}|{e['title']}\n")
    return True

def save_schedule(new_entry):
    """Додавання нового запису до txt файлу."""
    schedule = load_schedule()
    for date, preacher in new_entry.items():
        if date in schedule:
            if preacher not in schedule[date]:
                schedule[date].append(preacher)
        else:
            schedule[date] = [preacher]
    save_schedule_to_file(schedule)

async def export_table_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """
    Команда /export [current|next]
    Створює Word-документ з таблицею проповідників тільки за обраний місяць:
    - /export_table current => поточний місяць
    - /export_table next    => наступний місяць
    Якщо не передано аргумент, використовується поточний місяць.
    """
    if not is_admin_chat(update):
        return
    user_input = update.message.text.strip().split()
    chosen_option = None
    if len(user_input) > 1:
        chosen_option = user_input[1].lower()  # "current" або "next"

    # Визначаємо, який місяць фільтрувати
    now = datetime.now()
    this_year = now.year
    this_month = now.month

    if chosen_option == "next":
        # Наступний місяць
        if this_month == 12:
            filter_year = this_year + 1
            filter_month = 1
        else:
            filter_year = this_year
            filter_month = this_month + 1
    else:
        # За замовчуванням - поточний місяць
        filter_year = this_year
        filter_month = this_month

    # -------------------------------------------------
    # 1) Завантажуємо розклад з БД
    # -------------------------------------------------
    schedule = load_schedule()  
    if not schedule:
        await update.message.reply_text("Розклад порожній, немає що експортувати.")
        return

    # Усі дати, відсортовані
    all_dates_str = sorted(schedule.keys(), key=lambda d: datetime.strptime(d, "%d.%m.%Y"))
    if not all_dates_str:
        await update.message.reply_text("Немає дат у розкладі.")
        return

    # -------------------------------------------------
    # 2) Фільтруємо лише ті дати, які припадають на обраний місяць/рік
    # -------------------------------------------------
    filtered_dates = []
    for date_str in all_dates_str:
        dt = datetime.strptime(date_str, "%d.%m.%Y")
        if dt.year == filter_year and dt.month == filter_month:
            filtered_dates.append(date_str)

    if not filtered_dates:
        # Якщо немає дат за обраний місяць
        month_name = f"{filter_month:02d}.{filter_year}"
        await update.message.reply_text(
            f"Немає жодної дати для місяця {month_name} у розкладі."
        )
        return

    # -------------------------------------------------
    # 3) Cписок проповідників (може бути ваш PREACHERS = [...])
    # -------------------------------------------------
    preachers = [
        "Босько П.", "Біленко Ю.", "Мосійчук В.", "Мироненко І.",
        "Барановський М.", "Пономарьов А.", "Замуруєв В.", "Григоров А.",
        "Савостін В.", "Козак Є.", "Кулик Є.", "Ковальчук Ю.",
        "Савостін І.", "Суржа П.", "Волос В"
    ]

    # -------------------------------------------------
    # 4) Створюємо документ Word
    # -------------------------------------------------
    doc = docx.Document()

    # Щоб у підписі було зрозуміло, за який місяць генеруємо
    doc.add_heading(
        f"Розклад проповідей за {filter_month:02d}.{filter_year}",
        level=1
    )

    # Кількість рядків: 1 (шапка з датами) + кількість проповідників
    # Кількість колонок: 1 (список проповідників) + кількість дат
    rows_count = 1 + len(preachers)
    cols_count = 1 + len(filtered_dates)

    table = doc.add_table(rows=rows_count, cols=cols_count)
    table.style = "Table Grid"

    # -------------------------------------------------
    # 4.1) Заповнюємо верхній рядок (дата + день тижня)
    # -------------------------------------------------
    table.cell(0, 0).text = "Проповідники"

    short_days = ["Пн", "Вт", "Ср", "Чт", "Пт", "Сб", "Нд"]
    for col_idx, date_str in enumerate(filtered_dates, start=1):
        dt = datetime.strptime(date_str, "%d.%m.%Y")
        weekday_num = dt.weekday()  # 0=Пн ... 6=Нд
        day_of_week_str = short_days[weekday_num]
        cell = table.cell(0, col_idx)
        cell.text = f"{date_str}\n({day_of_week_str})"

    # -------------------------------------------------
    # 4.2) Заповнюємо перший стовпець проповідниками
    # -------------------------------------------------
    for row_idx, preacher in enumerate(preachers, start=1):
        table.cell(row_idx, 0).text = preacher

    # -------------------------------------------------
    # 4.3) Фарбуємо клітинки, якщо проповідник записаний на дату
    # -------------------------------------------------
    for col_idx, date_str in enumerate(filtered_dates, start=1):
        dt = datetime.strptime(date_str, "%d.%m.%Y")
        weekday_num = dt.weekday()

        for row_idx, preacher in enumerate(preachers, start=1):
            if date_str in schedule and preacher in schedule[date_str]:
                cell = table.cell(row_idx, col_idx)
                if weekday_num == 3:
                    # четвер => жовтий
                    set_cell_bg_color(cell, "FFFF00")
                elif weekday_num == 6:
                    # неділя => червоний
                    set_cell_bg_color(cell, "FF0000")
                else:
                    # Інші дні тижня — сірий (на бажання)
                    set_cell_bg_color(cell, "DDDDDD")

    # -------------------------------------------------
    # 5) Зберігаємо та надсилаємо файл
    # -------------------------------------------------
    filename = "графік.docx"
    doc.save(filename)

    with open(filename, "rb") as f:
        await context.bot.send_document(
            chat_id=update.effective_chat.id,
            document=f,
            filename=filename,
            caption=(
                f"Таблиця з розкладом проповідей за {filter_month:02d}.{filter_year}. "
                "Жовтий - четвер, червоний - неділя."
            )
        )

# Список проповідників
PREACHERS = [
    "Басько П.", "Біленко Ю.", "Мосійчук В.", "Мироненко І.",
    "Барановський М.", "Пономарьов А.", "Замуруєв В.", "Григоров А.",
    "Савостін В.", "Козак Є.", "Кулик Є.", "Ковальчук Ю.",
    "Савостін І.", "Суржа П.", "Кітченко Я.", "Волос В.",
    "Сардак Р."
]

# Мапа скорочень днів тижня
SHORT_DAYS_OF_WEEK = ["Пн", "Вт", "Ср", "Чт", "Пт", "Сб", "Нд"]

async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not is_admin_chat(update):
        return
    await update.message.reply_text("Привіт! Я готовий працювати. Використайте /help.")

async def help_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not is_admin_chat(update):
        return
    commands = """
Доступні команди:
/start - Почати спілкування з ботом
/add - Додати нову проповідь
/show - Показати розклад проповідей
/export - Експортувати розклад в файл
/delete - Видалити проповідь
/add_event - Додати церковну подію
/show_events - Показати заплановані події
/delete_event - Видалити подію
/help - Показати список доступних команд
"""
    await update.message.reply_text(commands)

def get_thursday_sunday_dates(year: int, month: int):
    """
    Повертає список дат (у форматі DD.MM.YYYY) для четвергів і неділь 
    у вказаному році та місяці.
    """
    # Дізнаємося, скільки днів у місяці
    _, days_in_month = calendar.monthrange(year, month)

    # Створюємо список усіх дат місяця
    all_dates = [
        datetime(year, month, day)
        for day in range(1, days_in_month + 1)
    ]

    # Потрібні лише четвер (weekday() = 3) та неділя (weekday() = 6)
    # (Пн=0, Вт=1, Ср=2, Чт=3, Пт=4, Сб=5, Нд=6)
    selected_dates = [
        d.strftime("%d.%m.%Y")
        for d in all_dates
        if d.weekday() in (3, 6)  # четвер і неділя
    ]

    return selected_dates

async def add_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not is_admin_chat(update):
        return
    now = datetime.now()
    current_year = now.year
    current_month = now.month

    # Дати четвергів і неділь для поточного місяця
    current_month_dates = get_thursday_sunday_dates(current_year, current_month)

    # Визначаємо наступний місяць
    if current_month == 12:
        next_year = current_year + 1
        next_month = 1
    else:
        next_year = current_year
        next_month = current_month + 1

    # Дати четвергів і неділь для наступного місяця
    next_month_dates = get_thursday_sunday_dates(next_year, next_month)

    # Об'єднуємо дати обох місяців (можна залишити окремо, але зручніше одним списком)
    all_dates = current_month_dates + next_month_dates

    if not all_dates:
        await update.message.reply_text(
            "Немає доступних четвергів чи неділь у поточному або наступному місяці."
        )
        return

    # Готуємо клавіатуру з датами
    keyboard = [[KeyboardButton(date_str)] for date_str in all_dates]

    # Записуємо стан користувача
    user_states[update.effective_user.id] = "waiting_for_date"

    # Виводимо повідомлення з вибором дати
    await update.message.reply_text(
        "Оберіть дату проповіді (лише четвер або неділя):",
        reply_markup=ReplyKeyboardMarkup(
            keyboard,
            one_time_keyboard=True,
            resize_keyboard=True
        ),
    )

def delete_schedule_date(date_str: str) -> bool:
    """Видаляє цілу дату (з усіма проповідниками) з txt файлу."""
    schedule = load_schedule()
    if date_str not in schedule:
        return False
    del schedule[date_str]
    save_schedule_to_file(schedule)
    return True

def delete_schedule_preacher(date_str: str, preacher: str) -> bool:
    """
    Видаляє вказаного проповідника з конкретної дати.
    Якщо після видалення проповідників не лишається — видаляє дату.
    """
    schedule = load_schedule()
    if date_str not in schedule:
        return False
    if preacher not in schedule[date_str]:
        return False
    schedule[date_str].remove(preacher)
    if not schedule[date_str]:
        del schedule[date_str]
    save_schedule_to_file(schedule)
    return True

async def delete_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not is_admin_chat(update):
        return
    schedule = load_schedule()
    if not schedule:
        await update.message.reply_text("Немає жодних дат у розкладі для видалення.")
        return

    # Отримуємо усі дати, що є в розкладі
    all_dates = sorted(schedule.keys())

    # Формуємо клавіатуру з дат
    keyboard = [[KeyboardButton(date_str)] for date_str in all_dates]

    # Ставимо стан для користувача
    user_states[update.effective_user.id] = "waiting_for_delete_date"

    await update.message.reply_text(
        "Оберіть дату, яку хочете видалити (цілком або окремих проповідників):",
        reply_markup=ReplyKeyboardMarkup(keyboard, one_time_keyboard=True, resize_keyboard=True),
    )

async def end_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not is_admin_chat(update):
        return
    schedule = load_schedule()
    if not schedule:
        await update.message.reply_text("Розклад порожній.")
        return

    result = "*Розклад проповідей:*\n\n"
    for date, preacher in sorted(schedule.items()):
        day_of_week = SHORT_DAYS_OF_WEEK[datetime.strptime(date, "%d.%m.%Y").weekday()]
        result += f"📆 {date} ({day_of_week}) Проповідники 🗣 {preacher}\n"

    await update.message.reply_text(result, parse_mode="Markdown")

async def get_chat_id(update: Update, context: ContextTypes.DEFAULT_TYPE):
    chat_id = update.effective_chat.id
    await update.message.reply_text(f"Chat ID цієї групи: {chat_id}")

async def error_handler(update: object, context: ContextTypes.DEFAULT_TYPE):
    print(f"Помилка: {context.error}")

async def handle_message(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = update.effective_user.id

    if user_id in user_states:
        state = user_states[user_id]

        # ---------------------------------------------------------------------
        #            СЦЕНАРІЙ ДОДАВАННЯ (ВЖЕ БУВ У ВАШОМУ КОДІ)
        # ---------------------------------------------------------------------
        if state == "waiting_for_date":
            selected_date = update.message.text.strip()
            user_states[user_id] = {
                "state": "waiting_for_preacher",
                "date": selected_date
            }
            keyboard = [[KeyboardButton(name)] for name in PREACHERS]
            await update.message.reply_text(
                "Оберіть проповідника:",
                reply_markup=ReplyKeyboardMarkup(
                    keyboard, one_time_keyboard=True, resize_keyboard=True
                ),
            )

        elif isinstance(state, dict) and state.get("state") == "waiting_for_preacher":
            preacher = update.message.text.strip()
            date = state["date"]

            # Зберігаємо запис у базі
            new_entry = {date: preacher}
            save_schedule(new_entry)

            schedule = load_schedule()
            propovidnyky = ", ".join(schedule[date])
            await update.message.reply_text(
                f"Проповідь на {date} збережено. Проповідники: {propovidnyky}"
            )

            # Очищаємо стан користувача
            del user_states[user_id]

        # ---------------------------------------------------------------------
        #            СЦЕНАРІЙ ВИДАЛЕННЯ
        # ---------------------------------------------------------------------
        elif state == "waiting_for_delete_date":
            # Користувач обрав дату для видалення
            chosen_date = update.message.text.strip()

            schedule = load_schedule()
            if chosen_date not in schedule:
                await update.message.reply_text("Такої дати в розкладі немає. Спробуйте ще раз.")
                return

            preachers = schedule[chosen_date]
            count = len(preachers)

            # Записуємо все необхідне в стан
            user_states[user_id] = {
                "state": "waiting_for_delete_decision",
                "date": chosen_date,
                "preachers": preachers
            }

            if count == 1:
                # Якщо лише один проповідник
                keyboard = [
                    [KeyboardButton("Видалити дату повністю")],
                    [KeyboardButton(f"Видалити проповідника: {preachers[0]}")],
                ]
                await update.message.reply_text(
                    f"Для дати {chosen_date} є лише один проповідник: {preachers[0]}.\n"
                    "Видалити дату цілком чи тільки проповідника?",
                    reply_markup=ReplyKeyboardMarkup(
                        keyboard, one_time_keyboard=True, resize_keyboard=True
                    ),
                )
            else:
                # Якщо проповідників декілька
                keyboard = [
                    [KeyboardButton("Видалити дату повністю")],
                    [KeyboardButton("Видалити одного проповідника")],
                ]
                await update.message.reply_text(
                    f"Для дати {chosen_date} є {count} проповідників: {', '.join(preachers)}.\n"
                    "Видалити дату цілком або лише одного?",
                    reply_markup=ReplyKeyboardMarkup(
                        keyboard, one_time_keyboard=True, resize_keyboard=True
                    ),
                )

        elif (
            isinstance(state, dict) and
            state.get("state") == "waiting_for_delete_decision"
        ):
            chosen_date = state["date"]
            preachers = state["preachers"]
            decision = update.message.text.strip().lower()

            if "повністю" in decision:
                # Користувач обрав видалити усю дату
                success = delete_schedule_date(chosen_date)
                if success:
                    await update.message.reply_text(
                        f"Дату {chosen_date} видалено повністю."
                    )
                else:
                    await update.message.reply_text(
                        f"Не вдалося видалити дату {chosen_date} (вона могла бути вже видалена)."
                    )
                del user_states[user_id]

            elif "одного" in decision:
                # Показуємо список проповідників для вибору
                user_states[user_id]["state"] = "waiting_for_delete_preacher"
                keyboard = [[KeyboardButton(p)] for p in preachers]
                await update.message.reply_text(
                    "Оберіть проповідника, якого хочете видалити:",
                    reply_markup=ReplyKeyboardMarkup(
                        keyboard, one_time_keyboard=True, resize_keyboard=True
                    ),
                )
            elif "проповідника:" in decision:
                # Це варіант, якщо один проповідник і кнопка мала вигляд "Видалити проповідника: Іванов І."
                preacher_to_delete = decision.split(":", 1)[1].strip()
                success = delete_schedule_preacher(chosen_date, preacher_to_delete)
                if success:
                    await update.message.reply_text(
                        f"Проповідника '{preacher_to_delete}' з дати {chosen_date} видалено."
                    )
                else:
                    await update.message.reply_text(
                        f"Не вдалося видалити проповідника '{preacher_to_delete}'."
                    )
                del user_states[user_id]

            else:
                # Невідомий варіант відповіді
                await update.message.reply_text("Невідома дія. Спробуйте ще раз /delete.")
                del user_states[user_id]

        elif (
            isinstance(state, dict) and
            state.get("state") == "waiting_for_delete_preacher"
        ):
            chosen_date = state["date"]
            chosen_preacher = update.message.text.strip()
            success = delete_schedule_preacher(chosen_date, chosen_preacher)
            if success:
                await update.message.reply_text(
                    f"Проповідника '{chosen_preacher}' з дати {chosen_date} видалено."
                )
            else:
                await update.message.reply_text(
                    f"Не вдалося видалити '{chosen_preacher}'. Можливо, немає такого проповідника."
                )
            del user_states[user_id]

        # ---------------------------------------------------------------------
        #            СЦЕНАРІЙ ДОДАВАННЯ ПОДІЇ
        # ---------------------------------------------------------------------
        elif state == "waiting_for_event_date":
            date_text = update.message.text.strip()
            try:
                datetime.strptime(date_text, "%d.%m.%Y")
            except ValueError:
                await update.message.reply_text(
                    "Невірний формат дати. Введіть у форматі ДД.ММ.РРРР (наприклад: 25.03.2026):"
                )
                return
            user_states[user_id] = {"state": "waiting_for_event_title", "date": date_text}
            await update.message.reply_text("Введіть назву події:")

        elif isinstance(state, dict) and state.get("state") == "waiting_for_event_title":
            title = update.message.text.strip()
            date = state["date"]
            save_event(date, title)
            await update.message.reply_text(f"Подію '{title}' на {date} збережено.")
            del user_states[user_id]

        # ---------------------------------------------------------------------
        #            СЦЕНАРІЙ ВИДАЛЕННЯ ПОДІЇ
        # ---------------------------------------------------------------------
        elif isinstance(state, dict) and state.get("state") == "waiting_for_delete_event":
            chosen = update.message.text.strip()
            events = state["events"]
            matched = next((e for e in events if f"{e['date']} — {e['title']}" == chosen), None)
            if not matched:
                await update.message.reply_text("Подію не знайдено. Спробуйте ще раз /delete_event.")
                del user_states[user_id]
                return
            success = delete_event(matched["date"], matched["title"])
            if success:
                await update.message.reply_text(f"Подію '{matched['title']}' на {matched['date']} видалено.")
            else:
                await update.message.reply_text("Не вдалося видалити подію.")
            del user_states[user_id]

async def add_event_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not is_admin_chat(update):
        return
    user_states[update.effective_user.id] = "waiting_for_event_date"
    await update.message.reply_text(
        "Введіть дату події у форматі ДД.ММ.РРРР (наприклад: 25.03.2026):"
    )

async def show_events_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not is_admin_chat(update):
        return
    events = load_events()
    now = datetime.now().date()
    upcoming = sorted(
        [e for e in events if datetime.strptime(e["date"], "%d.%m.%Y").date() >= now],
        key=lambda e: datetime.strptime(e["date"], "%d.%m.%Y")
    )
    if not upcoming:
        await update.message.reply_text("Немає запланованих подій.")
        return
    result = "*Заплановані події:*\n\n"
    for event in upcoming:
        dt = datetime.strptime(event["date"], "%d.%m.%Y")
        day_of_week = SHORT_DAYS_OF_WEEK[dt.weekday()]
        result += f"📅 {event['date']} ({day_of_week}) — {event['title']}\n"
    await update.message.reply_text(result, parse_mode="Markdown")

async def delete_event_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not is_admin_chat(update):
        return
    events = load_events()
    now = datetime.now().date()
    upcoming = sorted(
        [e for e in events if datetime.strptime(e["date"], "%d.%m.%Y").date() >= now],
        key=lambda e: datetime.strptime(e["date"], "%d.%m.%Y")
    )
    if not upcoming:
        await update.message.reply_text("Немає запланованих подій для видалення.")
        return
    keyboard = [[KeyboardButton(f"{e['date']} — {e['title']}")] for e in upcoming]
    user_states[update.effective_user.id] = {
        "state": "waiting_for_delete_event",
        "events": upcoming
    }
    await update.message.reply_text(
        "Оберіть подію для видалення:",
        reply_markup=ReplyKeyboardMarkup(keyboard, one_time_keyboard=True, resize_keyboard=True)
    )

async def remind(context: ContextTypes.DEFAULT_TYPE):
    try:
        current_date = datetime.now().date()

        # Нагадування про проповіді
        schedule = load_schedule()
        for date, preachers in schedule.items():
            schedule_date = datetime.strptime(date, "%d.%m.%Y").date()
            if (schedule_date - current_date).days == 2:
                preachers_list = ", ".join(preachers)
                await context.bot.send_message(
                    chat_id=GROUP_CHAT_ID,
                    text=(
                        f"Нагадування!\n\n"
                        f"На зібранні {date}:\n"
                        f"Проповідують: {preachers_list}"
                    )
                )

        # Нагадування про церковні події
        events = load_events()
        for event in events:
            event_date = datetime.strptime(event["date"], "%d.%m.%Y").date()
            if (event_date - current_date).days == 2:
                await context.bot.send_message(
                    chat_id=GROUP_CHAT_ID,
                    text=(
                        f"Нагадування про подію!\n\n"
                        f"📅 {event['date']}: {event['title']}"
                    )
                )

    except Exception as e:
        print(f"Помилка у функції remind: {e}")

# Обробник невідомої команди
async def unknown_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text(
        "Невідома команда. Використайте /help, щоб переглянути список доступних команд."
    )

def main():
    application = Application.builder().token(BOT_TOKEN).build()

    application.add_handler(CommandHandler("start", start))
    application.add_handler(CommandHandler("help", help_command))
    application.add_handler(CommandHandler("add", add_command))
    application.add_handler(CommandHandler("show", end_command))
    application.add_handler(CommandHandler("delete", delete_command))

    application.add_handler(CommandHandler("get_chat_id", get_chat_id))
    application.add_error_handler(error_handler)
    
    application.job_queue.run_repeating(remind, interval=24*60*60, first=10)
    # application.job_queue.run_repeating(remind, interval=10 )
    
    application.add_handler(CommandHandler("export", export_table_command))
    application.add_handler(CommandHandler("add_event", add_event_command))
    application.add_handler(CommandHandler("show_events", show_events_command))
    application.add_handler(CommandHandler("delete_event", delete_event_command))

    application.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, handle_message))
    application.add_handler(MessageHandler(filters.COMMAND, unknown_command))
  
    application.run_polling()

if __name__ == "__main__":
    main()
